from __future__ import annotations

import re
from io import BytesIO
from typing import Literal, Protocol

from pydantic import BaseModel, Field

from sharek_agents.agents.document_understanding.schemas import CloudinaryResourceRef


# ── Supported formats ──────────────────────────────────────────────────────────

SUPPORTED_FORMATS: frozenset[str] = frozenset({"pdf", "docx", "md", "txt"})

MIME_TO_FORMAT: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/markdown": "md",
    "text/x-markdown": "md",
    "text/plain": "txt",
}


# ── Error hierarchy ────────────────────────────────────────────────────────────


class ParsingError(Exception):
    """Base error for all document parsing failures."""


class UnsupportedFormatError(ParsingError):
    """The document format is not supported."""


class UnsupportedMimeTypeError(ParsingError):
    """The MIME type is not supported."""


class EmptyDocumentError(ParsingError):
    """The document has no content."""


class CorruptedDocumentError(ParsingError):
    """The document is corrupted and cannot be parsed."""


class EncryptedPDFError(ParsingError):
    """The PDF is encrypted and cannot be read without a password."""


class InvalidEncodingError(ParsingError):
    """The document has an invalid or unsupported encoding."""


class ParserLibraryError(ParsingError):
    """An error occurred in the underlying parsing library."""


# ── Parsed document models ─────────────────────────────────────────────────────


class DocumentElement(BaseModel):
    element_type: Literal["heading", "paragraph", "list_item", "code_block", "table"]
    text: str = Field(description="Extracted text content")
    level: int | None = Field(
        default=None, ge=1, description="Heading level (1=h1, 2=h2, ...)"
    )
    page_number: int | None = Field(
        default=None, ge=1, description="Page number within source document"
    )
    order: int = Field(
        default=0, ge=0, description="Zero-based position in document source order"
    )


class ParsedDocument(BaseModel):
    reference: CloudinaryResourceRef | None = Field(
        default=None, description="Original Cloudinary resource reference"
    )
    filename: str | None = Field(default=None, description="Source filename")
    file_format: str | None = Field(default=None, description="File format extension")
    content_type: str | None = Field(default=None, description="MIME type")
    elements: list[DocumentElement] = Field(
        default_factory=list,
        description="Ordered list of parsed document elements",
    )
    text: str = Field(
        default="", description="Full concatenated plain-text content"
    )
    page_count: int | None = Field(
        default=None, ge=0, description="Total number of pages if known"
    )
    metadata: dict = Field(
        default_factory=dict, description="Parser-specific metadata"
    )


# ── Parser Protocol ────────────────────────────────────────────────────────────


class DocumentParser(Protocol):
    """Protocol for parsers that convert raw document bytes into structured text.

    Each implementation handles one or more document formats.
    Parsers must not depend on LLMs, embeddings, vector stores, or network operations.
    """

    async def parse(
        self,
        content: bytes,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> ParsedDocument:
        """Parse raw bytes into a structured ``ParsedDocument``.

        Args:
            content: Raw document bytes.
            content_type: MIME type hint (e.g. ``application/pdf``).
            filename: Original filename (used for format detection fallback).

        Returns:
            A ``ParsedDocument`` with extracted elements and full text.

        Raises:
            EmptyDocumentError: The document has no content.
            CorruptedDocumentError: The document is corrupted.
            EncryptedPDFError: The PDF requires a password.
            InvalidEncodingError: The text encoding is unsupported.
            ParserLibraryError: The underlying parsing library failed.
            UnsupportedFormatError: The format is not supported.
        """
        ...


# ── PDF Parser ─────────────────────────────────────────────────────────────────


class PdfParser:
    """Parser for PDF documents using ``pypdf``.

    Extracts text page by page, preserving page numbers.
    Handles encrypted/corrupted PDFs safely.
    """

    async def parse(
        self,
        content: bytes,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> ParsedDocument:
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError
        except ImportError:
            raise ParserLibraryError(
                "pypdf is required for PDF parsing. Install with: pip install pypdf"
            )

        try:
            reader = PdfReader(BytesIO(content))
        except PdfReadError as exc:
            msg = str(exc).lower()
            if "encrypted" in msg or "password" in msg:
                raise EncryptedPDFError(
                    "The PDF is encrypted and cannot be read without a password"
                ) from exc
            raise CorruptedDocumentError(
                f"The PDF document is corrupted and cannot be parsed"
            ) from exc

        if reader.is_encrypted:
            raise EncryptedPDFError(
                "The PDF is encrypted and cannot be read without a password"
            )

        if len(reader.pages) == 0:
            raise EmptyDocumentError("The PDF document contains no pages")

        elements: list[DocumentElement] = []

        for i, page in enumerate(reader.pages):
            page_num = i + 1
            try:
                extracted = page.extract_text() or ""
            except Exception as exc:
                raise ParserLibraryError(
                    f"Failed to extract text from page {page_num}"
                ) from exc

            elements.append(
                DocumentElement(
                    element_type="paragraph",
                    text=extracted,
                    page_number=page_num,
                    order=i,
                )
            )

        text = "\n\n".join(e.text for e in elements)

        return ParsedDocument(
            elements=elements,
            text=text,
            page_count=len(reader.pages),
            metadata={"pdf_version": getattr(reader, "pdf_header", "")},
        )


# ── DOCX Parser ────────────────────────────────────────────────────────────────


class DocxParser:
    """Parser for DOCX documents using ``python-docx``.

    Extracts paragraphs, headings, lists, and tables in document order.
    Preserves heading hierarchy when detectable.
    """

    async def parse(
        self,
        content: bytes,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> ParsedDocument:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ParserLibraryError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )

        try:
            doc = Document(BytesIO(content))
        except Exception as exc:
            raise CorruptedDocumentError(
                "The DOCX document is invalid or corrupted and cannot be parsed"
            ) from exc

        if not doc.paragraphs and not doc.tables:
            raise EmptyDocumentError("The DOCX document contains no content")

        elements: list[DocumentElement] = []
        order = 0

        paras = list(doc.paragraphs)
        tables = list(doc.tables)
        para_idx = 0
        table_idx = 0

        body = doc.element.body

        for child in body:
            if child.tag == qn("w:p"):
                if para_idx >= len(paras):
                    para_idx += 1
                    continue
                para = paras[para_idx]
                para_idx += 1

                text = para.text
                if not text or not text.strip():
                    continue

                style_name = para.style.name if para.style else ""

                if style_name.startswith("Heading"):
                    parts = style_name.split()
                    level = 1
                    if len(parts) > 1 and parts[-1].isdigit():
                        level = int(parts[-1])
                    elements.append(
                        DocumentElement(
                            element_type="heading",
                            text=text.strip(),
                            level=level,
                            order=order,
                        )
                    )
                elif "List" in style_name or (
                    para._element.find(qn("w:numPr")) is not None
                ):
                    elements.append(
                        DocumentElement(
                            element_type="list_item",
                            text=text.strip(),
                            order=order,
                        )
                    )
                else:
                    elements.append(
                        DocumentElement(
                            element_type="paragraph",
                            text=text.strip(),
                            order=order,
                        )
                    )
                order += 1

            elif child.tag == qn("w:tbl"):
                if table_idx >= len(tables):
                    table_idx += 1
                    continue
                table = tables[table_idx]
                table_idx += 1

                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                table_text = "\n".join(rows_text) if rows_text else ""

                if table_text:
                    elements.append(
                        DocumentElement(
                            element_type="table",
                            text=table_text,
                            order=order,
                        )
                    )
                    order += 1

        full_text = "\n\n".join(e.text for e in elements)

        return ParsedDocument(elements=elements, text=full_text)


# ── Markdown Parser ────────────────────────────────────────────────────────────


class MarkdownParser:
    """Parser for Markdown (.md) documents.

    Preserves heading hierarchy, paragraphs, lists, code blocks,
    and tables with their source order.
    """

    async def parse(
        self,
        content: bytes,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> ParsedDocument:
        try:
            raw_text = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise InvalidEncodingError(
                "Markdown content must be valid UTF-8"
            ) from exc

        if not raw_text.strip():
            raise EmptyDocumentError("Markdown document is empty")

        elements: list[DocumentElement] = []
        order = 0

        in_code_block = False
        code_lines: list[str] = []
        lines = raw_text.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Code block fences
            if line.strip().startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_lines)
                    elements.append(
                        DocumentElement(
                            element_type="code_block",
                            text=code_text,
                            order=order,
                        )
                    )
                    order += 1
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                elements.append(
                    DocumentElement(
                        element_type="heading",
                        text=heading_text,
                        level=level,
                        order=order,
                    )
                )
                order += 1
                i += 1
                continue

            # Tables (consecutive | - delimited lines)
            if "|" in line and i + 1 < len(lines):
                next_line = lines[i + 1]
                if re.match(r"^[\s|:\-]+$", next_line):
                    table_lines: list[str] = []
                    while i < len(lines) and "|" in lines[i]:
                        table_lines.append(lines[i])
                        i += 1
                    table_text = "\n".join(table_lines)
                    elements.append(
                        DocumentElement(
                            element_type="table",
                            text=table_text,
                            order=order,
                        )
                    )
                    order += 1
                    continue

            # List items
            list_match = re.match(r"^(\s*)([-*+]|\d+[.)])\s+(.+)$", line)
            if list_match:
                list_text = list_match.group(3).strip()
                elements.append(
                    DocumentElement(
                        element_type="list_item",
                        text=list_text,
                        order=order,
                    )
                )
                order += 1
                i += 1
                continue

            # Paragraphs — collect consecutive non-empty, non-special lines
            para_lines: list[str] = []
            while i < len(lines):
                current = lines[i]
                stripped = current.strip()
                if not stripped:
                    break
                if re.match(r"^(#{1,6})\s", current):
                    break
                if re.match(r"^\s*[-*+] ", current) or re.match(
                    r"^\s*\d+[.)] ", current
                ):
                    break
                if current.strip().startswith("```"):
                    break
                if i + 1 < len(lines) and "|" in current:
                    next_candidate = lines[i + 1]
                    if re.match(r"^[\s|:\-]+$", next_candidate):
                        break
                para_lines.append(current)
                i += 1

            if para_lines:
                para_text = "\n".join(para_lines)
                elements.append(
                    DocumentElement(
                        element_type="paragraph",
                        text=para_text,
                        order=order,
                    )
                )
                order += 1

            while i < len(lines) and not lines[i].strip():
                i += 1

        full_text = "\n\n".join(e.text for e in elements)

        return ParsedDocument(elements=elements, text=full_text)


# ── Plain Text Parser ──────────────────────────────────────────────────────────


class TextParser:
    """Parser for plain text (.txt) documents.

    Preserves original line boundaries.
    Tries UTF-8 first, then falls back to common encodings.
    """

    FALLBACK_ENCODINGS: list[str] = ["utf-8", "utf-16", "latin-1", "cp1252"]

    async def parse(
        self,
        content: bytes,
        content_type: str | None = None,
        filename: str | None = None,
    ) -> ParsedDocument:
        raw_text: str | None = None
        last_error: Exception | None = None

        for encoding in self.FALLBACK_ENCODINGS:
            try:
                raw_text = content.decode(encoding)
                break
            except UnicodeDecodeError as exc:
                last_error = exc
                continue
            except LookupError:
                continue

        if raw_text is None:
            raise InvalidEncodingError(
                "Could not decode document content with any supported encoding"
            ) from last_error

        if not raw_text.strip():
            raise EmptyDocumentError("Text document is empty")

        elements: list[DocumentElement] = []
        order = 0

        for line in raw_text.split("\n"):
            elements.append(
                DocumentElement(
                    element_type="paragraph",
                    text=line,
                    order=order,
                )
            )
            order += 1

        return ParsedDocument(elements=elements, text="\n".join(raw_text.split("\n")))


# ── Format detection ────────────────────────────────────────────────────────────


def detect_format(
    file_format: str | None = None,
    content_type: str | None = None,
    filename: str | None = None,
) -> str | None:
    """Detect document format from available hints.

    Checks in priority order: ``file_format``, ``content_type``, ``filename``.
    Returns ``None`` when no supported format can be determined.
    """
    if file_format:
        fmt = file_format.lower().lstrip(".")
        if fmt in SUPPORTED_FORMATS:
            return fmt

    if content_type:
        fmt = MIME_TO_FORMAT.get(content_type.lower())
        if fmt:
            return fmt

    if filename:
        if "." in filename:
            ext = filename.rsplit(".", 1)[-1].lower()
            if ext in SUPPORTED_FORMATS:
                return ext

    return None


def detect_format_or_raise(
    file_format: str | None = None,
    content_type: str | None = None,
    filename: str | None = None,
) -> str:
    """Like ``detect_format`` but raises on failure."""
    fmt = detect_format(file_format, content_type, filename)
    if fmt is not None:
        return fmt

    if content_type:
        raise UnsupportedMimeTypeError(
            f"Unsupported MIME type: {content_type}. "
            f"Supported types: {', '.join(sorted(MIME_TO_FORMAT))}"
        )

    hint = file_format or "unknown"
    raise UnsupportedFormatError(
        f"Unsupported document format: {hint}. "
        f"Supported formats: {', '.join(sorted(SUPPORTED_FORMATS))}"
    )


# ── Parser registry ────────────────────────────────────────────────────────────

_PARSER_REGISTRY: dict[str, DocumentParser] = {
    "pdf": PdfParser(),
    "docx": DocxParser(),
    "md": MarkdownParser(),
    "txt": TextParser(),
}


def get_parser(file_format: str) -> DocumentParser:
    """Return the parser registered for *file_format*.

    Raises:
        UnsupportedFormatError: No parser is registered for this format.
    """
    fmt = file_format.lower().lstrip(".")
    parser = _PARSER_REGISTRY.get(fmt)
    if parser is None:
        raise UnsupportedFormatError(
            f"No parser available for format: {fmt}. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_FORMATS))}"
        )
    return parser


# ── Top-level parse function ───────────────────────────────────────────────────


async def parse_document(
    content: bytes,
    file_format: str | None = None,
    content_type: str | None = None,
    filename: str | None = None,
    reference: CloudinaryResourceRef | None = None,
) -> ParsedDocument:
    """Detect format, select a parser, and return a ``ParsedDocument``.

    Args:
        content: Raw document bytes.
        file_format: Explicit format hint (e.g. ``"pdf"``, ``"docx"``).
        content_type: MIME type hint.
        filename: Original filename (used as format detection fallback).
        reference: Optional ``CloudinaryResourceRef`` attached to the result.

    Returns:
        A ``ParsedDocument`` with source metadata and extracted elements.

    Raises:
        UnsupportedFormatError: No format could be detected.
        UnsupportedMimeTypeError: The MIME type is not supported.
        EmptyDocumentError: The document has no content.
        CorruptedDocumentError: The document is corrupted.
        EncryptedPDFError: The PDF requires a password.
        InvalidEncodingError: The text encoding is unsupported.
        ParserLibraryError: The underlying parsing library failed.
    """
    fmt = detect_format_or_raise(file_format, content_type, filename)
    parser = get_parser(fmt)

    result = await parser.parse(
        content=content,
        content_type=content_type,
        filename=filename,
    )

    result.reference = reference
    result.filename = filename or result.filename
    result.file_format = fmt
    if content_type and not result.content_type:
        result.content_type = content_type

    return result
